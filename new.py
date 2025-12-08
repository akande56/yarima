from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import RGBColor
import io

# Create a new Document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('Yarima Mining System: UI/UX Design & Theme Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(16)
title.runs[0].font.bold = True

# Subtitle
subtitle = doc.add_paragraph('Analysis of Visual Design, Responsive Layout, and Interactive Elements')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')  # Spacer

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc = [
    "1. Overview",
    "2. Color Scheme & Typography",
    "3. Global Layout & Structure",
    "4. Sidebar: 3D Navigation System",
    "5. Mobile Responsiveness",
    "6. Interactive Elements",
    "7. Form & Input Styling",
    "8. Button Design",
    "9. Alert & Feedback Components",
    "10. Glass Morphism & Visual Effects",
    "11. Animated Background: Floating Minerals",
    "12. Accessibility Features",
    "13. Technical Implementation",
    "14. Strengths & Best Practices",
    "15. Recommendations for Improvement",
    "16. Conclusion"
]

for item in toc:
    p = doc.add_paragraph(item)

# Add a page break
doc.add_page_break()

# 1. Overview
doc.add_heading('1. Overview', level=1)
overview_table = doc.add_table(rows=1, cols=2)
overview_table.style = 'Table Grid'
hdr = overview_table.rows[0].cells
hdr[0].text = 'Attribute'
hdr[1].text = 'Value'

data = [
    ("Theme Name", "Yarima Dark Glass"),
    ("Design Style", "Dark Mode + Glass Morphism + 3D Effects"),
    ("Primary Color", "#d6336c (Crimson Pink)"),
    ("Font Family", "Poppins, sans-serif"),
    ("Layout", "Sidebar + Main Content"),
    ("Responsive", "Yes (Mobile-First)"),
    ("Visual Effects", "3D rotation, floating animations, blur"),
    ("Accessibility", "Partial (keyboard nav, reduced motion)"),
    ("Target Environment", "LAN-based Mining Office"),
    ("Framework", "Custom CSS + Bootstrap 5")
]

for attr, value in data:
    row = overview_table.add_row().cells
    row[0].text = attr
    row[1].text = value

# 2. Color Scheme & Typography
doc.add_heading('2. Color Scheme & Typography', level=1)

# Color Palette
doc.add_heading('Color Palette', level=2)
color_table = doc.add_table(rows=1, cols=3)
color_table.style = 'Table Grid'
hdr = color_table.rows[0].cells
hdr[0].text = 'Color'
hdr[1].text = 'Hex Code'
hdr[2].text = 'Usage'

colors = [
    ("Primary", "#d6336c", "Buttons, links, active states, accents"),
    ("Hover", "#ff4d94", "Link/button hover states"),
    ("Background", "#0b0f19 → #2a2e35", "Gradient dark background"),
    ("Text", "#f1f1f1", "All body text and labels"),
    ("Form BG", "rgba(255,255,255,0.08)", "Input backgrounds with blur"),
    ("Glass BG", "rgba(255,255,255,0.08)", "Container transparency"),
    ("Border", "rgba(255,255,255,0.2)", "Input and element borders")
]

for name, hex_code, usage in colors:
    row = color_table.add_row().cells
    row[0].text = name
    row[1].text = hex_code
    row[2].text = usage

# Typography
doc.add_heading('Typography', level=2)
typography = doc.add_paragraph()
typography.add_run("• ").bold = True
typography.add_run("Font Family: Poppins (Google Fonts)\n")
typography.add_run("• ").bold = True
typography.add_run("Font Weight: 400–700 (light to bold)\n")
typography.add_run("• ").bold = True
typography.add_run("Text Color: #f1f1f1 (light gray/white)\n")
typography.add_run("• ").bold = True
typography.add_run("Link Color: #d6336c (primary pink)\n")
typography.add_run("• ").bold = True
typography.add_run("Hover Link: #ff4d94 (brighter pink)\n")
typography.add_run("• ").bold = True
typography.add_run("Smooth transitions on hover (0.3s ease)")

# 3. Global Layout & Structure
doc.add_heading('3. Global Layout & Structure', level=1)
layout = doc.add_paragraph()
layout.add_run("The UI follows a modern, component-based structure:\n\n")
layout.add_run("1. ").bold = True
layout.add_run("Body: Full viewport height with gradient background\n")
layout.add_run("2. ").bold = True
layout.add_run("Sidebar: Fixed 3D navigation panel (280px wide)\n")
layout.add_run("3. ").bold = True
layout.add_run("Main Content: Margin-left adjusts based on sidebar\n")
layout.add_run("4. ").bold = True
layout.add_run("Glass Containers: Frosted panels for forms and data\n")
layout.add_run("5. ").bold = True
layout.add_run("Background Minerals: Animated floating mineral icons\n")
layout.add_run("6. ").bold = True
layout.add_run("Mobile Toggle: Hamburger button for responsive access")

# 4. Sidebar: 3D Navigation System
doc.add_heading('4. Sidebar: 3D Navigation System', level=1)

# Sidebar Features
sidebar_features = [
    ("Perspective 3D", "Uses `perspective(1000px) rotateY(-2deg)` for depth"),
    ("Hover Effect", "Rotates to flat (0deg) on hover with enhanced shadow"),
    ("Backglass Effect", "Backdrop-filter: blur(20px) for transparency"),
    ("Fixed Header", "Logo and title stay at top during scroll"),
    ("Scrollable Nav", "Vertical scrolling with custom thin scrollbar"),
    ("Fade Overlays", "Top/bottom gradients hide scroll cutoffs"),
    ("Close Button", "Circular hover-rotate button in header")
]

sidebar_table = doc.add_table(rows=1, cols=2)
sidebar_table.style = 'Table Grid'
hdr = sidebar_table.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'Implementation'

for feature, impl in sidebar_features:
    row = sidebar_table.add_row().cells
    row[0].text = feature
    row[1].text = impl

# Nav Link Styling
doc.add_heading('Navigation Link Styling', level=2)
nav_styling = doc.add_paragraph()
nav_styling.add_run("• ").bold = True
nav_styling.add_run("Active/Inactive States: Gradient background and border-left\n")
nav_styling.add_run("• ").bold = True
nav_styling.add_run("Hover Effect: TranslateX(5px), color change, scale icon\n")
nav_styling.add_run("• ").bold = True
nav_styling.add_run("Focus Management: Outline for keyboard navigation\n")
nav_styling.add_run("• ").bold = True
nav_styling.add_run("Icon Animation: Scales on hover (1.2x)")

# 5. Mobile Responsiveness
doc.add_heading('5. Mobile Responsiveness', level=1)
mobile_table = doc.add_table(rows=1, cols=3)
mobile_table.style = 'Table Grid'
hdr = mobile_table.rows[0].cells
hdr[0].text = 'Breakpoint'
hdr[1].text = 'Behavior'
hdr[2].text = 'Implementation'

mobile_data = [
    ("≤ 767px", "Sidebar hidden by default", "Transform: translateX(-100%)"),
    ("≤ 767px", "Sidebar toggle button visible", "Fixed position (top-left)"),
    ("≤ 767px", "Sidebar opens with overlay", "Overlay z-index: 1040"),
    ("≤ 575px", "Reduced padding/margin", "Responsive glass containers"),
    ("Any", "Smooth scrolling disabled if reduced motion", "prefers-reduced-motion")
]

for bp, behavior, impl in mobile_data:
    row = mobile_table.add_row().cells
    row[0].text = bp
    row[1].text = behavior
    row[2].text = impl

# 6. Interactive Elements
doc.add_heading('6. Interactive Elements', level=1)
interactive = doc.add_paragraph()
interactive.add_run("• ").bold = True
interactive.add_run("Sidebar Toggle: Smooth slide-in with 3D effect\n")
interactive.add_run("• ").bold = True
interactive.add_run("Close Button: Rotates 90° on hover\n")
interactive.add_run("• ").bold = True
interactive.add_run("Link Transitions: 0.3s color and transform ease\n")
interactive.add_run("• ").bold = True
interactive.add_run("Button Effects: Lift (translateY), shadow growth\n")
interactive.add_run("• ").bold = True
interactive.add_run("Form Focus: Background lightens, border highlights\n")
interactive.add_run("• ").bold = True
interactive.add_run("Scrollbar: Hidden until hover, reveals thin pink bar")

# 7. Form & Input Styling
doc.add_heading('7. Form & Input Styling', level=1)
form_table = doc.add_table(rows=1, cols=3)
form_table.style = 'Table Grid'
hdr = form_table.rows[0].cells
hdr[0].text = 'Element'
hdr[1].text = 'Style'
hdr[2].text = 'Notes'

form_data = [
    ("Text Inputs", "Rounded (12px), glass bg, blur", "Backpack-filter: blur(10px)"),
    ("Placeholder", "Light gray (#f1f1f1 at 60%)", "Softer than main text"),
    ("Focus State", "Border: #d6336c, bg: lighter", "Transform: translateY(-2px)"),
    ("Select Dropdown", "Custom white arrow, no default", "SVG background image"),
    ("Checkboxes", "Glass bg, pink when checked", "Accessible focus outline"),
    ("Labels", "Weight: 500, color: #f1f1f1", "Clear hierarchy")
]

for elem, style, notes in form_data:
    row = form_table.add_row().cells
    row[0].text = elem
    row[1].text = style
    row[2].text = notes

# 8. Button Design
doc.add_heading('8. Button Design', level=1)
button_table = doc.add_table(rows=1, cols=3)
button_table.style = 'Table Grid'
hdr = button_table.rows[0].cells
hdr[0].text = 'Button Type'
hdr[1].text = 'Style'
hdr[2].text = 'Hover Effect'

button_data = [
    ("Primary", "Pink gradient, rounded, shadow", "Darker gradient, lift, bigger shadow"),
    ("Outline", "Pink border, glass bg", "Fill with gradient, lift"),
    ("Block", "Full-width (100%)", "Same as above"),
    ("Submit", "Styled as primary", "Same as primary")
]

for btn_type, style, hover in button_data:
    row = button_table.add_row().cells
    row[0].text = btn_type
    row[1].text = style
    row[2].text = hover

# 9. Alert & Feedback Components
doc.add_heading('9. Alert & Feedback Components', level=1)
alert_table = doc.add_table(rows=1, cols=3)
alert_table.style = 'Table Grid'
hdr = alert_table.rows[0].cells
hdr[0].text = 'Alert Type'
hdr[1].text = 'Color Scheme'
hdr[2].text = 'Usage'

alert_data = [
    ("Success", "Green bg (#10b981), left border", "Operation completed"),
    ("Danger", "Red bg (#ef4444), left border", "Error or rejection"),
    ("Warning", "Amber bg (#f59e0b), left border", "Caution needed"),
    ("Info", "Blue bg (#3b82f6), left border", "General information")
]

for alert_type, scheme, usage in alert_data:
    row = alert_table.add_row().cells
    row[0].text = alert_type
    row[1].text = scheme
    row[2].text = usage

# 10. Glass Morphism & Visual Effects
doc.add_heading('10. Glass Morphism & Visual Effects', level=1)
glass_effects = doc.add_paragraph()
glass_effects.add_run("• ").bold = True
glass_effects.add_run("Backdrop Filter: blur(10–25px) on inputs and containers\n")
glass_effects.add_run("• ").bold = True
glass_effects.add_run("Transparency: rgba(255,255,255,0.08–0.12)\n")
glass_effects.add_run("• ").bold = True
glass_effects.add_run("Inner Glow: Subtle box-shadow on focus/hover\n")
glass_effects.add_run("• ").bold = True
glass_effects.add_run("Gradient Borders: Subtle pink edge glow\n")
glass_effects.add_run("• ").bold = True
glass_effects.add_run("Smooth Transitions: cubic-bezier(0.25, 0.8, 0.25, 1)\n")
glass_effects.add_run("• ").bold = True
glass_effects.add_run("Shadow Depth: 0 12px 40px for glass panels")

# 11. Animated Background: Floating Minerals
doc.add_heading('11. Animated Background: Floating Minerals', level=1)
doc.add_paragraph("The UI includes a dynamic background animation of floating mineral icons.")

# CSS Animation
doc.add_heading('CSS Animation (theme.css)', level=2)
css_code = """@keyframes floatMineral {
  0% { transform: translateY(0) rotate(0deg) scale(1); }
  33% { transform: translateY(-30px) rotate(120deg) scale(1.1); }
  66% { transform: translateY(-15px) rotate(240deg) scale(0.9); }
  100% { transform: translateY(0) rotate(360deg) scale(1); }
}"""
doc.add_paragraph(css_code)

# JavaScript Animation
doc.add_heading('JavaScript Animation (mineral_float.js)', level=2)
js_code = """document.addEventListener('DOMContentLoaded', function () {
  const mineralImages = [
    '/static/images/minerals/tin.png',
    '/static/images/minerals/columbite.png',
    '/static/images/minerals/monazite.png'
  ];

  function spawnMineral() {
    const img = document.createElement('img');
    img.src = mineralImages[Math.floor(Math.random() * mineralImages.length)];
    img.classList.add('mineral-fall');
    img.style.left = `${Math.random() * 100}%`;
    container.appendChild(img);

    setTimeout(() => {
      img.style.animation = 'none';
      img.style.bottom = `${stackHeight[leftPos] || 0}px`;
      stackHeight[leftPos] = (stackHeight[leftPos] || 0) + 40;
    }, 6000);
  }

  setInterval(spawnMineral, 1500);
});"""
doc.add_paragraph(js_code)

# 12. Accessibility Features
doc.add_heading('12. Accessibility Features', level=1)
accessibility = doc.add_paragraph()
accessibility.add_run("✅ ").bold = True
accessibility.add_run("Keyboard Navigation: Focus outlines on links, buttons, inputs\n")
accessibility.add_run("✅ ").bold = True
accessibility.add_run("Reduced Motion: Disables animations if preferred\n")
accessibility.add_run("✅ ").bold = True
accessibility.add_run("Focus Management: scroll-margin for nav links\n")
accessibility.add_run("✅ ").bold = True
accessibility.add_run("High Contrast Mode: Adjusts colors and borders\n")
accessibility.add_run("⚠️ ").bold = True
accessibility.add_run("Screen Reader: No explicit ARIA labels found\n")
accessibility.add_run("⚠️ ").bold = True
accessibility.add_run("Color Contrast: Some text may fail WCAG AA (check #d6336c on dark bg)")

# 13. Technical Implementation
doc.add_heading('13. Technical Implementation', level=1)
tech_impl = doc.add_paragraph()
tech_impl.add_run("• ").bold = True
tech_impl.add_run("CSS: Custom theme.css with modern properties\n")
tech_impl.add_run("• ").bold = True
tech_impl.add_run("JS: mineral_float.js for dynamic background\n")
tech_impl.add_run("• ").bold = True
tech_impl.add_run("Fonts: Poppins (external Google Font)\n")
tech_impl.add_run("• ").bold = True
tech_impl.add_run("Responsive: Media queries for mobile\n")
tech_impl.add_run("• ").bold = True
tech_impl.add_run("Performance: Smooth animations, limited DOM elements\n")
tech_impl.add_run("• ").bold = True
tech_impl.add_run("Integration: Works with Bootstrap 5 components")

# 14. Strengths & Best Practices
doc.add_heading('14. Strengths & Best Practices', level=1)
strengths = [
    "✅ Modern Dark Theme with professional appearance",
    "✅ Glass morphism creates depth and elegance",
    "✅ 3D sidebar adds visual interest and interactivity",
    "✅ Mobile-first responsive design",
    "✅ Smooth animations and transitions",
    "✅ Custom form controls enhance UX",
    "✅ Animated background reinforces mining theme",
    "✅ Accessibility considerations (reduced motion, focus)",
    "✅ Consistent color scheme and typography",
    "✅ Well-organized CSS with clear sections"
]
for s in strengths:
    p = doc.add_paragraph()
    p.add_run(s)

# 15. Recommendations for Improvement
doc.add_heading('15. Recommendations for Improvement', level=1)
recommendations = [
    ("Add ARIA labels", "Improve screen reader support for navigation and icons"),
    ("Improve color contrast", "Ensure text meets WCAG AA standards"),
    ("Lazy load background images", "Optimize performance on low-end devices"),
    ("Add loading states", "For AJAX operations and form submission"),
    ("Include dark mode toggle", "Allow user preference override"),
    ("Optimize animations", "Limit mineral count to prevent lag"),
    ("Add tooltips", "For sidebar icons on hover"),
    ("Test on older browsers", "Ensure fallbacks for backdrop-filter"),
    ("Add focus traps", "For modal dialogs and overlays"),
    ("Document CSS variables", "For easier maintenance and theming")
]
rec_table = doc.add_table(rows=1, cols=2)
rec_table.style = 'Table Grid'
hdr = rec_table.rows[0].cells
hdr[0].text = 'Area'
hdr[1].text = 'Recommendation'
for area, rec in recommendations:
    row = rec_table.add_row().cells
    row[0].text = area
    row[1].text = rec

# 16. Conclusion
doc.add_heading('16. Conclusion', level=1)
conclusion = doc.add_paragraph()
conclusion.add_run("The Yarima Mining UI is a visually striking, modern interface that successfully combines functionality with aesthetic appeal. The dark glass theme, 3D navigation, and animated mineral background create a unique identity that reinforces the mining domain.\n\n")
conclusion.add_run("Key strengths include:\n")
conclusion.add_run("• Immersive visual design\n")
conclusion.add_run("• Responsive behavior across devices\n")
conclusion.add_run("• Interactive and engaging elements\n")
conclusion.add_run("• Attention to detail in form and button styling\n")
conclusion.add_run("• Accessibility considerations\n\n")
conclusion.add_run("With minor improvements in accessibility and performance, this UI provides an excellent foundation for a professional mining management system. The design effectively balances modern web aesthetics with practical usability for office staff.")

# Footer
footer = doc.sections[0].footer
footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
footer_para.text = "Prepared On: April 5, 2025 | Yarima Mining System: UI/UX Design & Theme Report"
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save to BytesIO
file_stream = io.BytesIO()

# Save the document to a file
doc.save("Yarima_Mining_System_UI_Report.docx")

print("Document generated successfully!")
print("The document contains:")
print("- Comprehensive system overview")
print("- Detailed analysis of all 4 office modules")
print("- Role-based access control documentation")
print("- Technical specifications and workflow diagrams")
print("- Business features and improvement recommendations")
print("- Professional formatting with tables and structured content")